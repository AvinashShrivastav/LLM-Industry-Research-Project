# streamlit_app.py
import streamlit as st
import os
import json
import re
import pandas as pd
from docx import Document
from backend1_2 import process_image_to_json, create_vector_db_from_pdf, generate_updated_json,plot_graph_from_json
# Importing the datetime module
from datetime import datetime
from docx.shared import Inches

fig = None
# Function to record inputs and outputs in a CSV file
def record_in_excel(image_name, original_json, updated_json):
    data = {'Image Name': [image_name],
            'Original JSON': [original_json],
            'Updated JSON': [updated_json]}
    df = pd.DataFrame(data)
    if not os.path.exists('record.csv'):
        df.to_csv('record.csv', index=False)
    else:
        df.to_csv('record.csv', mode='a', index=False, header=False)

# Function to generate a Word document
def generate_word_document(image_name, original_json, updated_json):
    output_dir = "output"
    document_path = os.path.join(output_dir, "output_document.docx")
    
    # Create the output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Check if the document already exists
    if os.path.exists(document_path):
        document = Document(document_path)
    else:
        document = Document()
        document.add_heading('Research Project Documentation', level=1)
        document.add_page_break()

    # Add date and time for the current input
    current_datetime = st.session_state.get('current_datetime', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    document.add_paragraph(f"Date and Time: {current_datetime}")
    
    # Add content for the current input
    document.add_heading(f'Input Image: {image_name}', level=2)
    original_image_path = os.path.join("uploads", image_name)
    document.add_picture(original_image_path, width=Inches(4))
    document.add_paragraph("Original JSON:", style='Caption')
    document.add_paragraph(json.dumps(original_json, indent=4))
    document.add_paragraph("Updated JSON:", style='Caption')
    document.add_paragraph(json.dumps(updated_json, indent=4))
    document.add_paragraph("Updated Chart:", style='Caption')
    plot_graph_from_json(updated_json, "updated_chart.png")
    updated_chart_path = "updated_chart.png"
    document.add_picture(updated_chart_path, width=Inches(4))
    document.add_page_break()

    # Save the document
    document.save(document_path)
    st.sidebar.success(f"Data saved to {document_path}")



def parse_json_string(json_string):

    # Remove <json> tags if present
    json_string = re.sub(r'^<json>\s*|\s*</json>$', '', json_string.strip())
    # Remove triple backticks if present
    json_string = re.sub(r'^```|```$', '', json_string.strip())
    
    try:
        json_object = json.loads(json_string)
    except json.JSONDecodeError as e:
        raise ValueError("Invalid JSON format") from e
    
    return json_object
# Main Streamlit app
def main():
    if 'current_datetime' not in st.session_state:
        st.session_state['current_datetime'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    st.set_page_config(page_title="Research Project", layout="wide")
    st.title("📈 Chart Updation Research App")
    st.markdown("Upload an image of a graph. The app will process the image to display original and updated JSONs, and the updated graph.")

    # Default PDF for vector database
    default_pdf_path = "uploads/chats.pdf"

    # Sidebar for file uploads
    st.sidebar.header("File Uploads")
    uploaded_image = st.sidebar.file_uploader("Upload an image of the chart", type=["jpg", "png", "jpeg"])
    uploaded_pdf = st.sidebar.file_uploader("Upload a PDF as vector DB", type=["pdf"])
    use_default_pdf = st.sidebar.checkbox("Use Default PDF for Vector DB")

    original_json = None
    vector_db_created = False
    original_image_path = None

    # Process uploaded image
    if uploaded_image is not None:
        original_image_path = os.path.join("uploads", uploaded_image.name)
        with open(original_image_path, "wb") as f:
            f.write(uploaded_image.getbuffer())

        st.sidebar.success("Image uploaded successfully!")

        # Process image to JSON
        original_json = parse_json_string(process_image_to_json(original_image_path))
        with st.spinner('Extracting Chart Data. Kindly wait...'):
            st.json(original_json)

    # Use uploaded or default PDF for vector DB
    if uploaded_pdf is not None or use_default_pdf:
        if uploaded_pdf is not None:
            pdf_path = os.path.join("uploads", uploaded_pdf.name)
            with open(pdf_path, "wb") as f:
                f.write(uploaded_pdf.getbuffer())
            vector_db_created = create_vector_db_from_pdf(pdf_path)
        else:
            vector_db_created = create_vector_db_from_pdf(default_pdf_path)

        if vector_db_created:
            st.sidebar.success("Vector DB created successfully!")
        else:
            st.sidebar.error("Failed to create Vector DB!")

    # Generate updated JSON and use sample image for updated graph
    if original_json and vector_db_created:
        with st.spinner('Updating json. Kindly wait...'):
            updated_json = generate_updated_json(original_json)

        # Record inputs and outputs in CSV file
        record_in_excel(os.path.basename(original_image_path), original_json, updated_json)

        # Generate Word document
        
        generate_word_document(os.path.basename(original_image_path), original_json, updated_json)

        # Display JSONs side by side
        col1, col2 = st.columns(2)
        with col1:
            st.write("### Original Graph JSON")
            st.json(original_json, expanded=True)
        with col2:
            st.write("### Updated Graph JSON")
            st.json(updated_json, expanded=True)

        # Display sample updated graph
        col1, col2 = st.columns(2)
        with col1:
            st.write("### Original Graph")
            st.image(original_image_path, caption="Sample Original Graph", use_column_width=True)
        with col2:
            st.write("### Updated Graph")
            updated_graph_image = "updated_chart.png"
            st.image(updated_graph_image, caption="Sample Updated Graph", use_column_width=True)
        

# Create upload and output directories if not exists
if __name__ == "__main__":
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    if not os.path.exists("output"):
        os.makedirs("output")
    main()


