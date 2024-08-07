import streamlit as st
import os
import json
import re
import pandas as pd
import time
from docx import Document
from docx.shared import Inches
from datetime import datetime
from backend3 import process_image_to_json, create_vector_db_from_pdf, generate_updated_json_gpt,generate_updated_json_gemini, plot_graph_from_json

fig = None
prompt_update = '''
You are an expert in data updation. You have got context and data to be updated.
Task is to understand the context of the data and update the data till 2024 with your own knowledge
and return JSON output'''
# Function to record inputs and outputs in a CSV file
def record_in_csv(image_name, original_json, updated_json_gpt,updated_json_gemini, prompt, processing_time,create_time_db ,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini):
    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data = {
        'Date and Time': [current_datetime],
        'Image Name': [image_name],
        'Original JSON': [json.dumps(original_json)],
        'Updated JSON GPT': [json.dumps(updated_json_gpt)],
        'Updated JSON Gemini ': [json.dumps(updated_json_gemini)],
        'Prompt Used': [prompt],
        'Processing Time (s)': [processing_time],
        'VB Creating Time (s)': [create_time_db],
        'Updating Time GPT(s)': [updating_time_gpt],
        'Updating Time Gemini(s)': [updating_time_gemini],
        'Plotting Time (s) GPT': [plotting_time_gpt],
        'Plotting Time (s) Gemini': [plotting_time_gemini],
        'Model Used (Image to JSON)': ['Gemini Vision Pro'],
        'Model Used (Update JSON)': ['ChatGPT 3.5 ang Gemini']
    }
    df = pd.DataFrame(data)
    if not os.path.exists('record_new.csv'):
        df.to_csv('record_new.csv', index=False)
    else:
        df.to_csv('record_new.csv', mode='a', index=False, header=False)

# Function to add a table to a Word document
def add_performance_table(document, processing_time,create_time_db ,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini):
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Task'
    hdr_cells[1].text = 'Model Used'
    hdr_cells[2].text = 'Description'
    hdr_cells[3].text = 'Time Taken (s)'

    tasks = [
        ('Processing Image to JSON', 'Gemini Vision Pro', 'Time taken to convert the uploaded image to JSON format', processing_time),
        ('Creating Vector DB', 'FAISS', 'Time taken to create vector DB', create_time_db),
        ('Updating JSON', 'GPT 3.5', 'Time taken to update the JSON with new data', updating_time_gpt),
        ('Updating JSON', 'Gemini', 'Time taken to update the JSON with new data', updating_time_gemini),
        ('Plotting Updated JSON', 'GPT 3.5', 'Time taken to plot the updated JSON', plotting_time_gpt),
        ('Plotting Updated JSON', 'Gemini', 'Time taken to plot the updated JSON', plotting_time_gemini),
    ]

    for task, model, description, time_taken in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = task
        row_cells[1].text = model
        row_cells[2].text = description
        row_cells[3].text = f"{time_taken:.2f}"

# Function to generate a Word document
def generate_word_document(image_name,  original_json, updated_json_gpt,updated_json_gemini, prompt, processing_time,create_time_db ,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini):
    output_dir = "output"
    document_path = os.path.join(output_dir, "output_document_new.docx")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    if os.path.exists(document_path):
        document = Document(document_path)
    else:
        document = Document()
        document.add_heading('Research Project Documentation', level=1)
        document.add_page_break()

    current_datetime = st.session_state.get('current_datetime', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    document.add_paragraph(f"Date and Time: {current_datetime}")
    
    document.add_heading(f'Input Image: {image_name}', level=2)
    original_image_path = os.path.join("uploads", image_name)
    document.add_picture(original_image_path, width=Inches(4))

    document.add_paragraph("Prompt Used:", style='Caption')
    document.add_paragraph(prompt, style='BodyText')

    document.add_paragraph("Original JSON:", style='Caption')
    document.add_paragraph(json.dumps(original_json, indent=4), style='BodyText')
    document.add_paragraph("Updated JSON GPT:", style='Caption')
    document.add_paragraph(json.dumps(updated_json_gpt, indent=4), style='BodyText')
    document.add_paragraph("Updated JSON Gemini:", style='Caption')
    document.add_paragraph(json.dumps(updated_json_gemini, indent=4), style='BodyText')
    document.add_paragraph("Updated Chart GPT:", style='Caption')
    document.add_picture("updated_chart_gpt.png", width=Inches(4))
    document.add_paragraph("Updated Chart Gemini:", style='Caption')
    document.add_picture("updated_chart_gemini.png", width=Inches(4))

    document.add_heading('Performance Metrics', level=2)
    add_performance_table(document, processing_time,create_time_db ,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini)

    document.add_page_break()
    document.save(document_path)
    st.sidebar.success(f"Data saved to {document_path}")


# Function to save prompt to file
def save_prompt(prompt_text, prompt_name):
    if not os.path.exists("prompts"):
        os.makedirs("prompts")
    with open(os.path.join("prompts", f"{prompt_name}.txt"), "w") as f:
        f.write(prompt_text)

# Function to load saved prompts from files
def load_saved_prompts():
    prompts = {}
    if os.path.exists("prompts"):
        for prompt_file in os.listdir("prompts"):
            with open(os.path.join("prompts", prompt_file), "r") as f:
                prompts[prompt_file] = f.read()
    return prompts

def main():
    if 'current_datetime' not in st.session_state:
        st.session_state['current_datetime'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    st.set_page_config(page_title="Research Project", layout="wide")
    st.title("📈 Chart Updation Research App")
    st.markdown("Upload an image of a graph and manage prompts for JSON updates.")

    default_pdf_path = "uploads/2024.pdf"
    st.sidebar.header("File Uploads")

    uploaded_image = st.sidebar.file_uploader("Upload an image of the chart", type=["jpg", "png", "jpeg"])
    uploaded_pdf = st.sidebar.file_uploader("Upload a PDF as vector DB", type=["pdf"])
    use_default_pdf = st.sidebar.checkbox("Use Default PDF for Vector DB")

    if uploaded_image is not None:
        st.session_state['uploaded_image'] = uploaded_image
        original_image_path = os.path.join("uploads", uploaded_image.name)
        with open(original_image_path, "wb") as f:
            f.write(uploaded_image.getbuffer())

        st.sidebar.success("Image uploaded successfully!")
        st.write("### Original Graph")
        st.image(original_image_path, caption="Original Graph", use_column_width=True)


    saved_prompts = load_saved_prompts()
    prompt_text = ""
    prompt_name = ""

    # Prompt selection
    st.markdown("### Select or Add Prompt")
    prompt_option = st.radio("Choose an option", ("Select a Saved Prompt", "Add a New Prompt"))

    if prompt_option == "Select a Saved Prompt":
        selected_prompt = st.selectbox("Select a Prompt", [""] + list(saved_prompts.keys()))
        if selected_prompt:
            prompt_text = saved_prompts[selected_prompt]
            st.text_area("Prompt", prompt_text, height=150, key="selected_prompt_text")

    elif prompt_option == "Add a New Prompt":
        prompt_text = st.text_area("Prompt", height=150, key="new_prompt_text")
        prompt_name = st.text_input("Prompt Name", key="new_prompt_name")

        if st.button("Save Prompt"):
            if prompt_text and prompt_name:
                save_prompt(prompt_text, prompt_name)
                st.success(f"Prompt saved as {prompt_name}.txt")
            else:
                st.error("Please enter both prompt text and name.")

    if 'uploaded_image' in st.session_state and uploaded_image:
        original_image_path = os.path.join("uploads", st.session_state['uploaded_image'].name)

        # Measure processing time
        start_time = time.time()
        original_json = process_image_to_json(original_image_path, prompt_text)
        processing_time = time.time() - start_time
        st.write("### Original Graph JSON")
        st.write(original_json, expanded=True)


        if original_json and st.button("Update Chart"):
            if use_default_pdf or uploaded_pdf:
                if use_default_pdf:
                    pdf_path = default_pdf_path
                else:
                    pdf_path = os.path.join("uploads", uploaded_pdf.name)
                with st.spinner("Creating Vector Data Base..."):
                    start_time = time.time()
                    vector_db = create_vector_db_from_pdf([pdf_path])
                    creat_time_db = time.time() - start_time

                if vector_db:
                    st.sidebar.success("Vector DB created successfully!")
                else:
                    st.sidebar.error("Failed to create Vector DB!")
                    return
                
                with st.spinner("Updating Chart from GPT..."):
                    start_time = time.time()
                    # updated_json_gpt = generate_updated_json_gpt(vector_db, original_json,prompt_update )
                    updated_json_gpt = original_json
                    updating_time_gpt = time.time() - start_time
                with st.spinner("Updating Chart from Gemini..."):
                    start_time = time.time() 
                    updated_json_gemini = generate_updated_json_gemini(vector_db,original_json,prompt_update )
                    updating_time_gemini = time.time() - start_time
                with st.spinner("PLotting Charts..."):
                    start_time = time.time()
                    plot_graph_from_json(updated_json_gpt, "updated_chart_gpt.png")
                    plotting_time_gpt = time.time() - start_time
                    start_time = time.time()
                    plot_graph_from_json(updated_json_gemini, "updated_chart_gemini.png")
                    plotting_time_gemini = time.time() - start_time

                # Save the prompt
                selected_prompt = st.session_state.get("selected_prompt_text", "")
                if prompt_option == "Select a Saved Prompt":
                    prompt_used = selected_prompt
                else:
                    prompt_used = st.session_state.get("new_prompt_text", "")

                # Record data
                record_in_csv(os.path.basename(original_image_path), original_json, updated_json_gpt,updated_json_gemini, prompt_used, processing_time, creat_time_db,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini)
                generate_word_document(os.path.basename(original_image_path), original_json, updated_json_gpt,updated_json_gemini, prompt_used, processing_time, creat_time_db,updating_time_gpt,updating_time_gemini, plotting_time_gpt, plotting_time_gemini)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.write("### Updated GPT Graph JSON")
                    st.json(updated_json_gpt, expanded=True)
                with col2:
                    st.write("### Updated Gemini Graph JSON")
                    st.json(updated_json_gemini, expanded=True)
                
                
                col1, col2 = st.columns(2)
                with col1:
                    st.write("### Updated Graph GPT")
                    st.image("updated_chart_gpt.png", caption="Updated gpt Graph", use_column_width=True)
                with col2:
                    st.write("### Updated Graph Gemini")
                    st.image("updated_chart_gemini.png", caption="Updated gemini Graph", use_column_width=True)

                # Display performance metrics
                st.write("### Performance Metrics")
                performance_data = {
                    "Task": ["Processing Image to JSON","Vector DB Creation", "Updating JSON","Updating JSON", "Plotting Updated JSON"],
                    "Model Used": ["Gemini Vision Pro","FAISS","ChatGPT 3.5", "Gemini pro","N/A"],
                    "Time Taken (s)": [f"{processing_time:.2f}",f"{creat_time_db:.2f}", f"{updating_time_gpt:.2f}",f"{updating_time_gemini:.2f}", f"{plotting_time:.2f}"]
                }
                performance_df = pd.DataFrame(performance_data)
                st.table(performance_df)
            else:
                st.error("Please select or upload a PDF for vector DB.")

    if st.sidebar.button("Clear"):
        st.session_state.clear()
        st.experimental_rerun()

if __name__ == "__main__":
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    if not os.path.exists("output"):
        os.makedirs("output")
    if not os.path.exists("prompts"):
        os.makedirs("prompts")
    main()
